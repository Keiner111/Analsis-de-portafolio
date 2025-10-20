# -*- coding: utf-8 -*-
import streamlit as st
import pandas as pd
from docx import Document
from datetime import datetime
from docx.shared import Inches # Para insertar imagenes en Word
import matplotlib.pyplot as plt # Para generar graficos
import io # Para manejar imagenes en memoria
import math # Para calculos matematicos como NaN e inf

# Importaciones para ReportLab (generacion de PDF)
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle # Importar ParagraphStyle
from reportlab.lib.units import inch # Para unidades de medida en PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter # Tamano de pagina (carta)

# Funcion para formatear valores monetarios en pesos colombianos
def formato_pesos(valor):
    # Formatea el numero con separador de miles como '.' y decimal como ','
    return f"${valor:,.0f}".replace(",", "X").replace(".", ",").replace("X", ".")

# Funcion para convertir un numero de meses (flotante) en meses enteros y dias
def convertir_a_meses_dias(meses):
    if meses == float('inf'):
        return "Mas de 50 anos", "" # Mensaje para tiempo infinito
    meses_int = int(meses)
    dias = round((meses - meses_int) * 30) # Asumiendo un promedio de 30 dias por mes
    return meses_int, dias

# Funcion principal para generar el informe completo (Word y PDF)
def generar_docx(df):
    st.markdown("<h1 style='text-align: center; color: #2E8B57;'>Generador de Informe Automatizado</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; font-size: 18px; color: #696969;'>Crea informes detallados de tu portafolio de inversiones.</p>", unsafe_allow_html=True)

    # Verifica si el DataFrame (df) esta vacio o no se ha cargado
    if df is None or df.empty:
        st.warning("No hay datos para generar el informe. Por favor, carga tu portafolio primero.")
        return

    # --- Parametros del Informe (Inflacion y Rentabilidad Anualizada) ---
    st.subheader("Parametros Clave del Informe")
    st.info("Estos parametros se usan para los calculos y proyecciones en tu informe.")

    # Inflacion anual: se toma de la variable de sesion establecida en el modulo devaluacion.py
    # Si el modulo devaluacion no ha sido visitado, se usa un valor por defecto.
    inflacion_anual_ejemplo = st.session_state.get('inflacion_anual_input', 5.05)


    # Rentabilidad anualizada: se toma de la variable de sesion.


    # Boton para iniciar la generacion de ambos informes
    if st.button("Generar informes (Word y PDF) "):
        st.spinner("Generando tu informe, por favor espera...")
        # --- Limpieza y preparacion de datos (Comun para DOCX y PDF) ---
        df_cleaned = df.copy()
        df_cleaned['Dinero'] = pd.to_numeric(df_cleaned['Dinero'].replace('[\$,]', '', regex=True), errors='coerce')
        df_cleaned['Interes Mensual'] = pd.to_numeric(df_cleaned['Interes Mensual'], errors='coerce').fillna(0)
        df_cleaned = df_cleaned.dropna(subset=['Dinero']) # Elimina filas con 'Dinero' no valido

        capital_total = df_cleaned['Dinero'].sum()
        ingreso_pasivo_mensual = df_cleaned['Interes Mensual'].sum()
        # Calcula la rentabilidad anualizada aproximada del portafolio
        rentabilidad_anual_portafolio = ((ingreso_pasivo_mensual * 12) / capital_total) * 100 if capital_total > 0 else 0
        rentabilidad_mensual_portafolio = (1 + rentabilidad_anual_portafolio / 100) ** (1/12) - 1

        # KPI: Porcentaje de capital productivo
        capital_productivo = df_cleaned[df_cleaned['Interes Mensual'] > 0]['Dinero'].sum()
        porcentaje_capital_productivo = (capital_productivo / capital_total * 100) if capital_total > 0 else 0

        # --- Obtener valores de la Ruta hacia la Meta desde st.session_state ---
        capital_meta_report = st.session_state.get('capital_meta_informe', 50_000_000.0)
        inversion_mensual_report = st.session_state.get('inversion_mensual_informe', 1_800_000.0)
        ingreso_pasivo_objetivo_report = st.session_state.get('ingreso_pasivo_objetivo_informe', 1_000_000.0)


        # --- Generacion de Graficos (Comun para DOCX y PDF) ---
        # Grafico de distribucion del portafolio por Tipo de inversion
        buf1 = None
        if 'Tipo de inversion' in df_cleaned.columns and not df_cleaned.empty:
            portfolio_distribution = df_cleaned.groupby('Tipo de inversion')['Dinero'].sum()
            if not portfolio_distribution.empty:
                fig1, ax1 = plt.subplots(figsize=(8, 8))
                ax1.pie(portfolio_distribution, labels=portfolio_distribution.index, autopct='%1.1f%%', startangle=90, pctdistance=0.85)
                ax1.axis('equal')
                ax1.set_title('Distribucion del Capital por Tipo de Inversion')
                
                buf1 = io.BytesIO()
                plt.savefig(buf1, format="png", bbox_inches='tight')
                plt.close(fig1)
                buf1.seek(0) # Rewind buffer for reuse

        # Grafico de ingresos pasivos por tipo de inversion
        buf2 = None
        if 'Tipo de inversion' in df_cleaned.columns and not df_cleaned.empty:
            income_by_type = df_cleaned.groupby('Tipo de inversion')['Interes Mensual'].sum()
            if not income_by_type.empty:
                fig2, ax2 = plt.subplots(figsize=(10, 6))
                income_by_type.plot(kind='bar', ax=ax2, color='skyblue')
                ax2.set_title('Ingresos Pasivos Mensuales por Tipo de Inversion')
                ax2.set_xlabel('Tipo de Inversion')
                ax2.set_ylabel('Ingreso Mensual (COP)')
                ax2.ticklabel_format(style='plain', axis='y')
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()

                buf2 = io.BytesIO()
                plt.savefig(buf2, format="png", bbox_inches='tight')
                plt.close(fig2)
                buf2.seek(0) # Rewind buffer for reuse

        # Calculate rendimiento_real_anual here, as it's used in Executive Summary
        rendimiento_real_anual = ( (1 + rentabilidad_anual_portafolio / 100) / (1 + inflacion_anual_ejemplo / 100) - 1 ) * 100

        # --- Generacion del Documento Word (DOCX) ---
        doc_docx = Document()
        doc_docx.add_heading('Informe Financiero Automatizado', 0)
        doc_docx.add_paragraph(f"Fecha de generacion: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

        # Nuevo: Resumen Ejecutivo (DOCX)
        doc_docx.add_page_break()
        doc_docx.add_heading('0. Resumen Ejecutivo', level=1)
        doc_docx.add_paragraph("Este informe proporciona un analisis detallado de tu portafolio de inversiones, proyectando tu camino hacia la independencia financiera y evaluando el impacto de factores economicos clave. A continuacion, los puntos mas destacados:")
        doc_docx.add_paragraph(f"- **Capital Total Consolidado:** {formato_pesos(capital_total)}")
        doc_docx.add_paragraph(f"- **Ingreso Pasivo Mensual Estimado:** {formato_pesos(ingreso_pasivo_mensual)}")
        doc_docx.add_paragraph(f"- **Rentabilidad Anual del Portafolio:** {rentabilidad_anual_portafolio:.2f}%")
        doc_docx.add_paragraph(f"- **Rendimiento Real Anual (ajustado por inflacion):** {rendimiento_real_anual:.2f}% {'(?Tu dinero esta creciendo en poder adquisitivo!)' if rendimiento_real_anual > 0 else '(Tu dinero esta perdiendo poder adquisitivo, se recomienda revision)' if rendimiento_real_anual < 0 else '(Tu dinero mantiene su poder adquisitivo)'}")
        doc_docx.add_paragraph(f"- **Porcentaje de Capital Productivo:** {porcentaje_capital_productivo:.2f}%")
        
        meses_i, dias_i = convertir_a_meses_dias(0) # Default for display
        if capital_total < capital_meta_report:
            temp_capital_con_interes = capital_total
            meses_con_interes_resumen = 0
            while temp_capital_con_interes < capital_meta_report and meses_con_interes_resumen <= 600:
                temp_capital_con_interes = (temp_capital_con_interes + inversion_mensual_report) * (1 + rentabilidad_mensual_portafolio)
                meses_con_interes_resumen += 1
            if temp_capital_con_interes < capital_meta_report:
                meses_i_resumen, dias_i_resumen = "Mas de 50 anos", ""
            else:
                meses_i_resumen, dias_i_resumen = convertir_a_meses_dias(meses_con_interes_resumen)
        else:
            meses_i_resumen, dias_i_resumen = 0, 0 # Already reached goal
            
        doc_docx.add_paragraph(f"- **Tiempo Estimado para Meta ({formato_pesos(capital_meta_report)} con interes compuesto):** {meses_i_resumen} meses y {dias_i_resumen} dias (aprox. {meses_i_resumen // 12} anos y {meses_i_resumen % 12} meses)")
        doc_docx.add_paragraph(f"- **Capital FIRE Estimado:** {formato_pesos(capital_fire_estimado if 'capital_fire_estimado' in locals() else 0)}") # Ensure capital_fire_estimado is defined, though it will be in full run
        if capital_total >= (capital_fire_estimado if 'capital_fire_estimado' in locals() else float('inf')):
            doc_docx.add_paragraph("- **Estado FIRE:** ?Felicidades! Has alcanzado o superado tu capital FIRE estimado.")
        else:
            doc_docx.add_paragraph("- **Estado FIRE:** Necesitas acumular mas capital productivo para alcanzar tu meta FIRE.")
        doc_docx.add_paragraph("Este resumen ofrece una instantanea de tu salud financiera actual. Para un analisis detallado, consulta las secciones siguientes.")
        

        # 1. Analisis del Portafolio (DOCX)
        doc_docx.add_page_break()
        doc_docx.add_heading('1. Analisis Detallado del Portafolio', level=1)
        doc_docx.add_paragraph(f"**Capital total consolidado:** {formato_pesos(capital_total)}. Este es el valor acumulado de todas tus inversiones.")
        doc_docx.add_paragraph(f"**Ingreso pasivo mensual estimado:** {formato_pesos(ingreso_pasivo_mensual)}. Representa los ingresos recurrentes que generas de tus inversiones cada mes, sin requerir trabajo activo.")
        doc_docx.add_paragraph(f"**Rentabilidad anual aproximada del portafolio:** {rentabilidad_anual_portafolio:.2f}%. Esta es la tasa de retorno que tu portafolio ha generado anualmente.")
        doc_docx.add_paragraph(f"**Porcentaje de capital productivo:** {porcentaje_capital_productivo:.2f}%. Este KPI indica que proporcion de tu capital esta generando ingresos pasivos. Un porcentaje mas alto sugiere mayor eficiencia en tu portafolio.")

        doc_docx.add_heading('Distribucion del Capital por Tipo de Inversion', level=2)
        if buf1:
            doc_docx.add_picture(buf1, width=Inches(6))
            buf1.seek(0) # Rewind for potential PDF reuse if not already done
            doc_docx.add_paragraph("El grafico muestra la composicion de tu portafolio, destacando como se distribuye tu capital entre diferentes tipos de inversion. Una diversificacion adecuada es clave para mitigar riesgos.")
        else:
            doc_docx.add_paragraph("No hay datos suficientes para generar el grafico de distribucion del portafolio. Asegurate de tener al menos una inversion cargada.")

        doc_docx.add_heading('Ingresos Pasivos por Tipo de Inversion', level=2)
        if buf2:
            doc_docx.add_picture(buf2, width=Inches(6))
            buf2.seek(0) # Rewind for potential PDF reuse if not already done
            doc_docx.add_paragraph("Este grafico ilustra que tipos de inversion son tus principales fuentes de ingreso pasivo. Identificar estas fuentes te ayuda a optimizar y fortalecer tus flujos de efectivo.")
        else:
            doc_docx.add_paragraph("No hay datos suficientes para generar el grafico de ingresos pasivos. Asegurate de que tus inversiones generen intereses mensuales.")


        # 2. Ruta hacia la Meta Financiera (DOCX)
        doc_docx.add_page_break()
        doc_docx.add_heading('2. Proyeccion y Ruta Hacia tu Meta Financiera', level=1)
        
        meses_con_interes = 0
        temp_capital_con_interes = capital_total
        while temp_capital_con_interes < capital_meta_report and meses_con_interes <= 600:
            temp_capital_con_interes = (temp_capital_con_interes + inversion_mensual_report) * (1 + rentabilidad_mensual_portafolio)
            meses_con_interes += 1
        if temp_capital_con_interes < capital_meta_report:
            meses_con_interes = float('inf')

        meses_sin_interes = 0.0
        if inversion_mensual_report > 0:
            if capital_total >= capital_meta_report:
                meses_sin_interes = 0.0
            else:
                meses_sin_interes = (capital_meta_report - capital_total) / inversion_mensual_report
        else:
            meses_sin_interes = float('inf')

        doc_docx.add_paragraph(f"**Capital meta proyectado:** {formato_pesos(capital_meta_report)}. Este es el monto de capital que has establecido como objetivo a alcanzar.")
        doc_docx.add_paragraph(f"**Aporte mensual adicional considerado:** {formato_pesos(inversion_mensual_report)}. Este es el valor que planeas invertir adicionalmente cada mes para acelerar el crecimiento de tu capital.")
        
        meses_i, dias_i = convertir_a_meses_dias(meses_con_interes)
        doc_docx.add_paragraph(f"**Tiempo estimado para alcanzar la meta (con interes compuesto):** {meses_i} meses y {dias_i} dias (~{meses_i // 12} anos y {meses_i % 12} meses). Esta proyeccion considera el efecto multiplicador de tus rendimientos reinvertidos.")

        if meses_sin_interes == float('inf'):
            doc_docx.add_paragraph(f"**Tiempo estimado para alcanzar la meta (sin considerar interes):** No se puede calcular (aporte mensual = 0). Para lograr tu meta sin ingresos pasivos, es esencial un aporte mensual significativo.")
        else:
            meses_s, dias_s = convertir_a_meses_dias(meses_sin_interes)
            doc_docx.add_paragraph(f"**Tiempo estimado para alcanzar la meta (sin considerar interes):** {meses_s} meses y {dias_s} dias (~{meses_s // 12} anos y {meses_s % 12} meses). Esta simulacion muestra el tiempo que tomaria alcanzar tu meta solo con tus aportes, sin el beneficio del interes compuesto.")
        
        doc_docx.add_paragraph("Este analisis resalta la **potencia del interes compuesto** en la aceleracion de tus objetivos financieros. Cada ganancia reinvertida contribuye exponencialmente a tu crecimiento patrimonial.")


        # 3. Analisis FIRE (DOCX)
        doc_docx.add_page_break()
        doc_docx.add_heading('3. Analisis de Independencia Financiera (FIRE)', level=1)
        tasa_retiro_segura = 4
        gastos_anuales_estimados = ingreso_pasivo_mensual * 12
        capital_fire_estimado = (gastos_anuales_estimados / tasa_retiro_segura) * 100 if tasa_retiro_segura > 0 else float('inf')

        doc_docx.add_paragraph(f"**Ingresos pasivos anuales actuales:** {formato_pesos(ingreso_pasivo_mensual * 12)}. Este es el total de tus ingresos pasivos proyectados a un ano.")
        doc_docx.add_paragraph(f"**Capital FIRE estimado (Regla del {tasa_retiro_segura}%):** {formato_pesos(capital_fire_estimado)}. Este valor representa el capital necesario para que tus ingresos pasivos puedan cubrir tus gastos anuales, asumiendo una tasa de retiro del 4%.")
        doc_docx.add_paragraph(f"**Estado actual frente a la meta FIRE:**")
        if capital_total >= capital_fire_estimado:
            doc_docx.add_paragraph("?Felicidades! Tu capital actual es **suficiente para alcanzar la independencia financiera** segun la regla del 4%. Esto significa que tus ingresos pasivos podrian cubrir tus gastos anuales, permitiendote la libertad de elegir si trabajar o no.")
        else:
            doc_docx.add_paragraph(f"Para **solidificar tu posicion FIRE**, necesitas acumular aproximadamente **{formato_pesos(capital_fire_estimado - capital_total)}** adicional en capital productivo. Este es un objetivo clave para tu libertad economica y te acerca a la capacidad de vivir de tus inversiones.")

        # 4. Impacto de la Inflacion y Rentabilidad Anualizada (DOCX)
        doc_docx.add_page_break()
        doc_docx.add_heading('4. Impacto de la Inflacion y Rentabilidad Anualizada en tu Patrimonio', level=1)
        # Usar la rentabilidad anual del portafolio calculado en la seccion anterior
        
        doc_docx.add_paragraph(f"**Inflacion anual estimada:** {inflacion_anual_ejemplo:.2f}%. La inflacion es el aumento general de los precios y la perdida del poder adquisitivo de la moneda.")
        # Ahora se toma la rentabilidad calculada del portafolio
        doc_docx.add_paragraph(f"**Rentabilidad anualizada estimada del portafolio:** {rentabilidad_anual_portafolio:.2f}%. Esta es la tasa de crecimiento anual de tu capital invertido.")
        doc_docx.add_paragraph(f"**Rendimiento real anual de tu portafolio (ajustado por inflacion):** {rendimiento_real_anual:.2f}%.")
        if rendimiento_real_anual > 0:
            doc_docx.add_paragraph("Este valor es **crucial** porque indica cuanto crece tu **poder adquisitivo real** despues de descontar el efecto de la inflacion. Un rendimiento real positivo significa que tu dinero esta ganando valor con el tiempo, permitiendote comprar mas bienes y servicios en el futuro.")
        elif rendimiento_real_anual < 0:
            doc_docx.add_paragraph("Este valor es **crucial** porque indica cuanto estas perdiendo en **poder adquisitivo real** despues de descontar el efecto de la inflacion. Un rendimiento real negativo significa que tu dinero esta perdiendo valor con el tiempo, lo que implica que podras comprar menos bienes y servicios en el futuro. Es fundamental revisar tu estrategia de inversion para superar la inflacion.")
        else:
            doc_docx.add_paragraph("Este valor es **crucial** porque indica que tu **poder adquisitivo real** se mantiene estable despues de descontar el efecto de la inflacion. Tu dinero no esta ganando ni perdiendo valor real, lo cual es mejor que perder, pero aun hay oportunidades para un crecimiento real.")

        # 5. Sugerencias de Rebalanceo (DOCX)
        doc_docx.add_page_break()
        doc_docx.add_heading('5. Estrategias para el Rebalanceo de tu Portafolio', level=1)
        doc_docx.add_paragraph("El rebalanceo es una **practica fundamental** para mantener tu portafolio **alineado con tus objetivos de riesgo y retorno** a lo largo del tiempo. Permite **optimizar la asignacion de activos** y **mitigar la exposicion a riesgos no deseados**.")
        doc_docx.add_paragraph("Basado en un analisis general de tu portafolio, te proponemos las siguientes consideraciones estrategicas:")
        doc_docx.add_paragraph("- **Considera la posibilidad de incrementar tu exposicion en activos de renta fija** si tu perfil de riesgo tiende a ser mas conservador, buscando asi una **mayor estabilidad y previsibilidad** en tus rendimientos.")
        doc_docx.add_paragraph("- Para un perfil mas dinamico, **explora activamente nuevas oportunidades de inversion** en sectores con **alto potencial de crecimiento** o en mercados emergentes, lo que podria **potenciar tus retornos a largo plazo**.")
        doc_docx.add_paragraph("- Es **imperativo revisar periodicamente tus asignaciones de activos** para asegurar que se **ajusten continuamente a tus metas financieras** y a las **condiciones cambiantes del mercado**. Un ajuste proactivo puede **maximizar tus ganancias** y **minimizar perdidas**.")

        # 6. Evaluacion de Activos Fisicos (DOCX)
        doc_docx.add_page_break()
        doc_docx.add_heading('6. Evaluacion y Optimizacion de Activos Fisicos', level=1)
        doc_docx.add_paragraph("Los activos fisicos, como bienes raices o semovientes, constituyen una **parte significativa de tu patrimonio**, ofreciendo **diversificacion y potencial de valorizacion**. Su gestion adecuada es **esencial para la salud financiera** de tu portafolio.")
        doc_docx.add_paragraph("En tu portafolio, se han identificado activos como:")
        
        activos_fisicos_df = df_cleaned[df_cleaned['Tipo de inversion'].isin(['Animal- semoviente', 'Activo Fisico'])]
        if not activos_fisicos_df.empty:
            for _, row in activos_fisicos_df.iterrows():
                doc_docx.add_paragraph(f"- **{row['Items']}**: Valor actual de {formato_pesos(row['Dinero'])}. Este activo puede ofrecer {'' if row['Interes Mensual'] == 0 else 'ingresos pasivos adicionales o'} una proteccion contra la inflacion, aunque su liquidez puede ser menor.")
        else:
            doc_docx.add_paragraph("- No se identificaron activos fisicos especificos en tu portafolio cargado que cumplan con la clasificacion de 'Animal- semoviente' o 'Activo Fisico'.")

        doc_docx.add_paragraph("Es **altamente recomendable evaluar periodicamente el rendimiento y la liquidez** de estos activos, asi como su **contribucion efectiva a la diversificacion general** de tu patrimonio. La **optimizacion de su gestion** puede **desbloquear un valor adicional** y **mejorar la eficiencia** de tu capital.")

        # 7. Inversiones Detalladas del Portafolio (DOCX)
        doc_docx.add_page_break()
        doc_docx.add_heading('7. Detalle Exhaustivo de las Inversiones del Portafolio', level=1)
        doc_docx.add_paragraph("A continuacion, se presenta una tabla con el detalle completo de cada una de las inversiones registradas en tu portafolio. Esta informacion es fundamental para una **comprension granular** de tu exposicion y rendimiento.")
        tabla_docx = doc_docx.add_table(rows=1, cols=len(df.columns))
        tabla_docx.style = 'Table Grid'
        hdr_cells_docx = tabla_docx.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells_docx[i].text = col
        for _, row in df.iterrows():
            fila_docx = tabla_docx.add_row().cells
            for i, val in enumerate(row):
                fila_docx[i].text = str(val)

        # 8. Recomendaciones Generales Adicionales (DOCX)
        doc_docx.add_page_break()
        doc_docx.add_heading("8. Recomendaciones Estrategicas Adicionales", level=1)
        doc_docx.add_paragraph("Para **fortalecer aun mas tu posicion financiera** y **acelerar la consecucion de tus objetivos**, te ofrecemos las siguientes sugerencias estrategicas:")
        doc_docx.add_paragraph("1. **Potenciar Inversiones de Alto Rendimiento:** Es **altamente beneficioso** identificar y **aumentar la asignacion de capital** en aquellas inversiones que consistentemente te estan generando los mayores ingresos pasivos. Esta estrategia puede **acelerar significativamente el crecimiento** de tu patrimonio.")
        doc_docx.add_paragraph("2. **Reevaluar Activos Suboptimos:** Te **aconsejamos encarecidamente revisar** aquellas inversiones que son improductivas o que presentan tasas de rendimiento inferiores al 0.5%. **Explorar alternativas mas rentables** o **redireccionar esos fondos** puede **mejorar la eficiencia general** de tu portafolio.")
        doc_docx.add_paragraph("3. **Aprovechar el Poder del Interes Compuesto:** **Considera la reinversion sistematica** de tus ingresos pasivos. Este habito puede **multiplicar tus ganancias exponencialmente** y **acortar dramaticamente tu camino** hacia la meta financiera.")
        doc_docx.add_paragraph("4. **Consolidar un Fondo de Emergencia Robusto:** Es **fundamental asegurar** un colchon financiero adecuado para imprevistos. Idealmente, este fondo deberia **cubrir entre 3 y 6 meses de tus gastos esenciales**, proporcionandote **tranquilidad y seguridad** ante cualquier eventualidad.")
        doc_docx.add_paragraph("5. **Implementar una Diversificacion Inteligente:** Para **mitigar riesgos y potenciar retornos**, es **crucial no concentrar** todos tus recursos en un solo tipo de activo. Una **diversificacion bien estructurada** a traves de diferentes clases de activos, sectores y geografias puede **blindar tu portafolio** contra la volatilidad del mercado.")
        doc_docx.add_paragraph("6. **Cultivar la Educacion Financiera Continua:** Mantenerte **informado y actualizado** sobre las nuevas oportunidades de inversion, las tendencias del mercado y las estrategias financieras emergentes es **esencial para tomar decisiones informadas** y **adaptarte a un entorno economico dinamico**.")


        # --- Generacion del Documento PDF ---
        pdf_buffer = io.BytesIO()
        doc_pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter)
        styles = getSampleStyleSheet()

        # Estilos personalizados para ReportLab (negrita, etc.)
        styles.add(ParagraphStyle(name='NormalBold', parent=styles['Normal'], fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='H1Bold', parent=styles['h1'], fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='H2Bold', parent=styles['h2'], fontName='Helvetica-Bold'))

        elements = []
        elements.append(Paragraph("<b>Informe Financiero Automatizado</b>", styles['h1']))
        elements.append(Paragraph(f"Fecha de generacion: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
        elements.append(Spacer(1, 0.2 * inch))

        # Nuevo: Resumen Ejecutivo (PDF)
        elements.append(PageBreak())
        elements.append(Paragraph("<b>0. Resumen Ejecutivo</b>", styles['H1Bold']))
        elements.append(Paragraph("Este informe proporciona un analisis detallado de tu portafolio de inversiones, proyectando tu camino hacia la independencia financiera y evaluando el impacto de factores economicos clave. A continuacion, los puntos mas destacados:", styles['Normal']))
        elements.append(Paragraph(f"- <b>Capital Total Consolidado:</b> {formato_pesos(capital_total)}", styles['NormalBold']))
        elements.append(Paragraph(f"- <b>Ingreso Pasivo Mensual Estimado:</b> {formato_pesos(ingreso_pasivo_mensual)}", styles['NormalBold']))
        elements.append(Paragraph(f"- <b>Rentabilidad Anual del Portafolio:</b> {rentabilidad_anual_portafolio:.2f}%", styles['NormalBold']))
        elements.append(Paragraph(f"- <b>Rendimiento Real Anual (ajustado por inflacion):</b> {rendimiento_real_anual:.2f}% {'(?Tu dinero esta creciendo en poder adquisitivo!)' if rendimiento_real_anual > 0 else '(Tu dinero esta perdiendo poder adquisitivo, se recomienda revision)' if rendimiento_real_anual < 0 else '(Tu dinero mantiene su poder adquisitivo)'}", styles['NormalBold']))
        elements.append(Paragraph(f"- <b>Porcentaje de Capital Productivo:</b> {porcentaje_capital_productivo:.2f}%", styles['NormalBold']))

        if capital_total < capital_meta_report:
            if meses_con_interes_resumen == "Mas de 50 anos":
                elements.append(Paragraph(f"- <b>Tiempo Estimado para Meta ({formato_pesos(capital_meta_report)} con interes compuesto):</b> {meses_i_resumen}", styles['NormalBold']))
            else:
                elements.append(Paragraph(f"- <b>Tiempo Estimado para Meta ({formato_pesos(capital_meta_report)} con interes compuesto):</b> {meses_i_resumen} meses y {dias_i_resumen} dias (aprox. {meses_i_resumen // 12} anos y {meses_i_resumen % 12} meses)", styles['NormalBold']))
        else:
            elements.append(Paragraph(f"- <b>Tiempo Estimado para Meta ({formato_pesos(capital_meta_report)} con interes compuesto):</b> ?Meta ya alcanzada!", styles['NormalBold']))


        elements.append(Paragraph(f"- <b>Capital FIRE Estimado:</b> {formato_pesos(capital_fire_estimado)}", styles['NormalBold']))
        if capital_total >= capital_fire_estimado:
            elements.append(Paragraph("- <b>Estado FIRE:</b> ?Felicidades! Has alcanzado o superado tu capital FIRE estimado.", styles['Normal']))
        else:
            elements.append(Paragraph("- <b>Estado FIRE:</b> Necesitas acumular mas capital productivo para alcanzar tu meta FIRE.", styles['Normal']))
        elements.append(Paragraph("Este resumen ofrece una instantanea de tu salud financiera actual. Para un analisis detallado, consulta las secciones siguientes.", styles['Normal']))
        elements.append(PageBreak())

        # 1. Analisis del Portafolio (PDF)
        elements.append(Paragraph("<b>1. Analisis Detallado del Portafolio</b>", styles['H1Bold']))
        elements.append(Paragraph(f"<b>Capital total consolidado:</b> {formato_pesos(capital_total)}. Este es el valor acumulado de todas tus inversiones.", styles['NormalBold']))
        elements.append(Paragraph(f"<b>Ingreso pasivo mensual estimado:</b> {formato_pesos(ingreso_pasivo_mensual)}. Representa los ingresos recurrentes que generas de tus inversiones cada mes, sin requerir trabajo activo.", styles['NormalBold']))
        elements.append(Paragraph(f"<b>Rentabilidad anual aproximada del portafolio:</b> {rentabilidad_anual_portafolio:.2f}%. Esta es la tasa de retorno que tu portafolio ha generado anualmente.", styles['NormalBold']))
        elements.append(Paragraph(f"<b>Porcentaje de capital productivo:</b> {porcentaje_capital_productivo:.2f}%. Este KPI indica que proporcion de tu capital esta generando ingresos pasivos. Un porcentaje mas alto sugiere mayor eficiencia en tu portafolio.", styles['NormalBold']))
        elements.append(Spacer(1, 0.2 * inch))

        elements.append(Paragraph("<b>Distribucion del Capital por Tipo de Inversion</b>", styles['H2Bold']))
        if buf1:
            if buf1.tell() != 0:
                buf1.seek(0)
            elements.append(Image(buf1, width=4*inch, height=4*inch)) # Ajusta tamano para PDF
            buf1.seek(0)
            elements.append(Paragraph("El grafico muestra la composicion de tu portafolio, destacando como se distribuye tu capital entre diferentes tipos de inversion. Una diversificacion adecuada es clave para mitigar riesgos.", styles['Normal']))
        else:
            elements.append(Paragraph("No hay datos suficientes para generar el grafico de distribucion del portafolio. Asegurate de tener al menos una inversion cargada.", styles['Normal']))
        elements.append(Spacer(1, 0.2 * inch))

        elements.append(Paragraph("<b>Ingresos Pasivos por Tipo de Inversion</b>", styles['H2Bold']))
        if buf2:
            if buf2.tell() != 0:
                buf2.seek(0)
            elements.append(Image(buf2, width=5*inch, height=3*inch)) # Ajusta tamano para PDF
            buf2.seek(0)
            elements.append(Paragraph("Este grafico ilustra que tipos de inversion son tus principales fuentes de ingreso pasivo. Identificar estas fuentes te ayuda a optimizar y fortalecer tus flujos de efectivo.", styles['Normal']))
        else:
            elements.append(Paragraph("No hay datos suficientes para generar el grafico de ingresos pasivos. Asegurate de que tus inversiones generen intereses mensuales.", styles['Normal']))
        elements.append(Spacer(1, 0.2 * inch))
        elements.append(PageBreak())

        # 2. Ruta hacia la Meta Financiera (PDF)
        elements.append(Paragraph("<b>2. Proyeccion y Ruta Hacia tu Meta Financiera</b>", styles['H1Bold']))
        
        meses_i, dias_i = convertir_a_meses_dias(meses_con_interes)
        elements.append(Paragraph(f"<b>Capital meta proyectado:</b> {formato_pesos(capital_meta_report)}. Este es el monto de capital que has establecido como objetivo a alcanzar.", styles['NormalBold']))
        elements.append(Paragraph(f"<b>Aporte mensual adicional considerado:</b> {formato_pesos(inversion_mensual_report)}. Este es el valor que planeas invertir adicionalmente cada mes para acelerar el crecimiento de tu capital.", styles['NormalBold']))
        elements.append(Paragraph(f"<b>Tiempo estimado para alcanzar la meta (con interes compuesto):</b> {meses_i} meses y {dias_i} dias (~{meses_i // 12} anos y {meses_i % 12} meses). Esta proyeccion considera el efecto multiplicador de tus rendimientos reinvertidos.", styles['NormalBold']))
        if meses_sin_interes == float('inf'):
            elements.append(Paragraph(f"<b>Tiempo estimado para alcanzar la meta (sin considerar interes):</b> No se puede calcular (aporte mensual = 0). Para lograr tu meta sin ingresos pasivos, es esencial un aporte mensual significativo.", styles['NormalBold']))
        else:
            meses_s, dias_s = convertir_a_meses_dias(meses_sin_interes)
            elements.append(Paragraph(f"<b>Tiempo estimado para alcanzar la meta (sin considerar interes):</b> {meses_s} meses y {dias_s} dias (~{meses_s // 12} anos y {meses_s % 12} meses). Esta simulacion muestra el tiempo que tomaria alcanzar tu meta solo con tus aportes, sin el beneficio del interes compuesto.", styles['NormalBold']))
        elements.append(Paragraph("Este analisis resalta la <b>potencia del interes compuesto</b> en la aceleracion de tus objetivos financieros. Cada ganancia reinvertida contribuye exponencialmente a tu crecimiento patrimonial.", styles['Normal']))
        elements.append(PageBreak())

        # 3. Analisis FIRE (PDF)
        elements.append(Paragraph("<b>3. Analisis de Independencia Financiera (FIRE)</b>", styles['H1Bold']))
        elements.append(Paragraph(f"<b>Ingresos pasivos anuales actuales:</b> {formato_pesos(ingreso_pasivo_mensual * 12)}. Este es el total de tus ingresos pasivos proyectados a un ano.", styles['NormalBold']))
        elements.append(Paragraph(f"<b>Capital FIRE estimado (Regla del {tasa_retiro_segura}%):</b> {formato_pesos(capital_fire_estimado)}. Este valor representa el capital necesario para que tus ingresos pasivos puedan cubrir tus gastos anuales, asumiendo una tasa de retiro del 4%.", styles['NormalBold']))
        elements.append(Paragraph(f"<b>Estado actual frente a la meta FIRE:</b>", styles['NormalBold']))
        if capital_total >= capital_fire_estimado:
            elements.append(Paragraph("?Felicidades! Tu capital actual es <b>suficiente para alcanzar la independencia financiera</b> segun la regla del 4%. Esto significa que tus ingresos pasivos podrian cubrir tus gastos anuales, permitiendote la libertad de elegir si trabajar o no.", styles['Normal']))
        else:
            elements.append(Paragraph("- <b>Estado FIRE:</b> Necesitas acumular mas capital productivo para alcanzar tu meta FIRE.", styles['Normal']))
        elements.append(PageBreak())

        # 4. Impacto de la Inflacion y Rentabilidad Anualizada (PDF)
        elements.append(Paragraph("<b>4. Impacto de la Inflacion y Rentabilidad Anualizada en tu Patrimonio</b>", styles['H1Bold']))
        elements.append(Paragraph(f"<b>Inflacion anual estimada:</b> {inflacion_anual_ejemplo:.2f}%. La inflacion es el aumento general de los precios y la perdida del poder adquisitivo de la moneda.", styles['NormalBold']))
        elements.append(Paragraph(f"<b>Rentabilidad anualizada estimada del portafolio:</b> {rentabilidad_anual_portafolio:.2f}%. Esta es la tasa de crecimiento anual de tu capital invertido.", styles['NormalBold']))
        elements.append(Paragraph(f"<b>Rendimiento real anual de tu portafolio (ajustado por inflacion):</b> {rendimiento_real_anual:.2f}%.", styles['NormalBold']))
        if rendimiento_real_anual > 0:
            elements.append(Paragraph("Este valor es <b>crucial</b> porque indica cuanto crece tu <b>poder adquisitivo real</b> despues de descontar el efecto de la inflacion. Un rendimiento real positivo significa que tu dinero esta ganando valor con el tiempo, permitiendote comprar mas bienes y servicios en el futuro.", styles['Normal']))
        elif rendimiento_real_anual < 0:
            elements.append(Paragraph("Este valor es <b>crucial</b> porque indica cuanto estas perdiendo en <b>poder adquisitivo real</b> despues de descontar el efecto de la inflacion. Un rendimiento real negativo significa que tu dinero esta perdiendo valor con el tiempo, lo que implica que podras comprar menos bienes y servicios en el futuro. Es fundamental revisar tu estrategia de inversion para superar la inflacion.", styles['Normal']))
        else:
            elements.append(Paragraph("Este valor es <b>crucial</b> porque indica que tu <b>poder adquisitivo real</b> se mantiene estable despues de descontar el efecto de la inflacion. Tu dinero no esta ganando ni perdiendo valor real, lo cual es mejor que perder, pero aun hay oportunidades para un crecimiento real.", styles['Normal']))
        elements.append(PageBreak())

        # 5. Sugerencias de Rebalanceo (PDF)
        elements.append(Paragraph("<b>5. Estrategias para el Rebalanceo de tu Portafolio</b>", styles['H1Bold']))
        elements.append(Paragraph("El rebalanceo es una <b>practica fundamental</b> para mantener tu portafolio <b>alineado con tus objetivos de riesgo y retorno</b> a lo largo del tiempo. Permite <b>optimizar la asignacion de activos</b> y <b>mitigar la exposicion a riesgos no deseados</b>.", styles['Normal']))
        elements.append(Paragraph("Basado en un analisis general de tu portafolio, te proponemos las siguientes consideraciones estrategicas:", styles['Normal']))
        elements.append(Paragraph("- <b>Considera la posibilidad de incrementar tu exposicion en activos de renta fija</b> si tu perfil de riesgo tiende a ser mas conservador, buscando asi una <b>mayor estabilidad y previsibilidad</b> en tus rendimientos.", styles['Normal']))
        elements.append(Paragraph("- Para un perfil mas dinamico, <b>explora activamente nuevas oportunidades de inversion</b> en sectores con <b>alto potencial de crecimiento</b> o en mercados emergentes, lo que podria <b>potenciar tus retornos a largo plazo</b>.", styles['Normal']))
        elements.append(Paragraph("- Es <b>imperativo revisar periodicamente tus asignaciones de activos</b> para asegurar que se <b>ajusten continuamente a tus metas financieras</b> y a las <b>condiciones cambiantes del mercado</b>. Un ajuste proactivo puede <b>maximizar tus ganancias</b> y <b>minimizar perdidas</b>.", styles['Normal']))
        elements.append(PageBreak())

        # 6. Evaluacion de Activos Fisicos (PDF)
        elements.append(Paragraph("<b>6. Evaluacion y Optimizacion de Activos Fisicos</b>", styles['H1Bold']))
        elements.append(Paragraph("Los activos fisicos, como bienes raices o semovientes, constituyen una <b>parte significativa de tu patrimonio</b>, ofreciendo <b>diversificacion y potencial de valorizacion</b>. Su gestion adecuada es <b>esencial para la salud financiera</b> de tu portafolio.", styles['Normal']))
        elements.append(Paragraph("En tu portafolio, se han identificado activos como:", styles['Normal']))
        activos_fisicos_df = df_cleaned[df_cleaned['Tipo de inversion'].isin(['Animal- semoviente', 'Activo Fisico'])]
        if not activos_fisicos_df.empty:
            for _, row in activos_fisicos_df.iterrows():
                elements.append(Paragraph(f"- <b>{row['Items']}</b>: Valor actual de {formato_pesos(row['Dinero'])}. Este activo puede ofrecer {'' if row['Interes Mensual'] == 0 else 'ingresos pasivos adicionales o'} una proteccion contra la inflacion, aunque su liquidez puede ser menor.", styles['Normal']))
        else:
            elements.append(Paragraph("- No se identificaron activos fisicos especificos en tu portafolio cargado que cumplan con la clasificacion de 'Animal- semoviente' o 'Activo Fisico'.", styles['Normal']))
        elements.append(Paragraph("Es <b>altamente recomendable evaluar periodicamente el rendimiento y la liquidez</b> de estos activos, asi como su <b>contribucion efectiva a la diversificacion general</b> de tu patrimonio. La <b>optimizacion de su gestion</b> puede <b>desbloquear un valor adicional</b> y <b>mejorar la eficiencia</b> de tu capital.", styles['Normal']))
        elements.append(PageBreak())

        # 7. Inversiones Detalladas del Portafolio (PDF)
        elements.append(Paragraph("<b>7. Detalle Exhaustivo de las Inversiones del Portafolio</b>", styles['H1Bold']))
        # Prepara los datos para la tabla de ReportLab
        table_data_pdf = [df.columns.tolist()] + df.astype(str).values.tolist()
        table_style_pdf = TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black)
        ])
        elements.append(Table(table_data_pdf, style=table_style_pdf))
        elements.append(PageBreak())

        # 8. Recomendaciones Generales Adicionales (PDF)
        elements.append(Paragraph("<b>8. Recomendaciones Estrategicas Adicionales</b>", styles['H1Bold']))
        elements.append(Paragraph("Para <b>fortalecer aun mas tu posicion financiera</b> y <b>acelerar la consecucion de tus objetivos</b>, te ofrecemos las siguientes sugerencias estrategicas:", styles['Normal']))
        elements.append(Paragraph("1. <b>Potenciar Inversiones de Alto Rendimiento:</b> Es <b>altamente beneficioso</b> identificar y <b>aumentar la asignacion de capital</b> en aquellas inversiones que consistentemente te estan generando los mayores ingresos pasivos. Esta estrategia puede <b>acelerar significativamente el crecimiento</b> de tu patrimonio.", styles['Normal']))
        elements.append(Paragraph("2. <b>Reevaluar Activos Suboptimos:</b> Te <b>aconsejamos encarecidamente revisar</b> aquellas inversiones que son improductivas o que presentan tasas de rendimiento inferiores al 0.5%. <b>Explorar alternativas mas rentables</b> o <b>redireccionar esos fondos</b> puede <b>mejorar la eficiencia general</b> de tu portafolio.", styles['Normal']))
        elements.append(Paragraph("3. <b>Aprovechar el Poder del Interes Compuesto:</b> <b>Considera la reinversion sistematica</b> de tus ingresos pasivos. Este habito puede <b>multiplicar tus ganancias exponencialmente</b> y <b>acortar dramaticamente tu camino</b> hacia la meta financiera.", styles['Normal']))
        elements.append(Paragraph("4. <b>Consolidar un Fondo de Emergencia Robusto:</b> Es <b>fundamental asegurar</b> un colchon financiero adecuado para imprevistos. Idealmente, este fondo deberia <b>cubrir entre 3 y 6 meses de tus gastos esenciales</b>, proporcionandote <b>tranquilidad y seguridad</b> ante cualquier eventualidad.", styles['Normal']))
        elements.append(Paragraph("5. <b>Implementar una Diversificacion Inteligente:</b> Para <b>mitigar riesgos y potenciar retornos</b>, es <b>crucial no concentrar</b> todos tus recursos en un solo tipo de activo. Una <b>diversificacion bien estructurada</b> a traves de diferentes clases de activos, sectores y geografias puede <b>blindar tu portafolio</b> contra la volatilidad del mercado.", styles['Normal'])) # <-- Se agreg¨® el par¨¦ntesis de cierre aqu¨ª.
        elements.append(Paragraph("6. <b>Cultivar la Educacion Financiera Continua:</b> Mantenerte <b>informado y actualizado</b> sobre las nuevas oportunidades de inversion, las tendencias del mercado y las estrategias financieras emergentes es <b>esencial para tomar decisiones informadas</b> y <b>adaptarte a un entorno economico dinamico</b>.", styles['Normal']))
        elements.append(PageBreak())

        # Guarda los informes en buffers
        docx_buffer = io.BytesIO()
        doc_docx.save(docx_buffer)
        docx_buffer.seek(0)

        # Build PDF
        doc_pdf.build(elements)
        pdf_buffer.seek(0)

        st.success("?Informes generados exitosamente! ")
       

        # Opciones de descarga
        st.download_button(
            label="Descargar Informe Word (.docx) ",
            data=docx_buffer,
            file_name=f"informe_financiero_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.download_button(
            label="Descargar Informe PDF (.pdf) ",
            data=pdf_buffer,
            file_name=f"informe_financiero_{datetime.now().strftime('%Y-%m-%d_%H%M%S')}.pdf",
            mime="application/pdf"
        )